import streamlit as st
import hashlib
from pymongo import MongoClient
from datetime import datetime
from bson.objectid import ObjectId
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import pandas as pd
import urllib.parse
import uuid
import pytz

import base64

st.set_page_config(page_title="Sistema Escolar - CCMPP by Leandro Malheiros V2.0.3 ", layout="centered")

# --- Estilização Visual ---
st.markdown("""
    <style>
        .stApp {
            background-color: #f2f6fc;
            color: #333333;
            font-family: 'Segoe UI', sans-serif;
        }
        h1, h2, h3 {
            color: #003366;
        }
        .block-container {
            max-width: 1000px;
            margin: auto;
            padding: 2rem;
            background-color: white;
            border-radius: 15px;
            box-shadow: 2px 2px 15px rgba(0,0,0,0.1);
        }
        .css-1d391kg {
            padding: 2rem 1rem 2rem 1rem;
        }
        .stButton>button {
            background-color: #003366;
            color: white;
            border-radius: 8px;
            padding: 0.5rem 1rem;
            font-weight: bold;
        }
        .stButton>button:hover {
            background-color: #0055a5;
            transition: 0.3s;
        }
        .stSelectbox, .stTextInput, .stTextArea {
            background-color: #e9f0fa;
            border-radius: 8px;
        }
        .stMarkdown {
            font-size: 1.1rem;
        }
    </style>
""", unsafe_allow_html=True)

def agora_local():
    tz = pytz.timezone("America/Sao_Paulo")
    return datetime.now(tz)
    
# --- Conexão com MongoDB ---
@st.cache_resource
def conectar():
    uri = "mongodb+srv://lsouzamalheiros_db_user:Malheiros76@cluster0.dmp6gyn.mongodb.net/"
    cliente = MongoClient(uri)
    return cliente["escola"]

db = conectar()
# --- Criar usuário admin padrão (primeira execução) ---
def criar_admin_padrao():
    if db.usuarios.count_documents({}) == 0:
        usuario_padrao = "admin"
        senha_padrao = "admin123"

        senha_hash = hashlib.sha256(senha_padrao.encode()).hexdigest()

        db.usuarios.insert_one({
            "usuario": usuario_padrao,
            "senha": senha_hash,
            "nivel": "admin"
        })

        print("✅ Usuário admin padrão criado!")
        try:
            st.warning("⚠️ Usuário padrão criado → admin / admin123 (altere após login!)")
        except:
            pass

# CHAMADA DA FUNÇÃO
criar_admin_padrao()

print("--- Coleções no banco 'escola' ---")
print(db.list_collection_names())

# --- Funções auxiliares ---
from datetime import datetime
import pandas as pd

def data_segura(valor):
    try:
        if not valor:
            return agora_local().date()

        data = pd.to_datetime(valor, errors="coerce")

        if pd.isna(data):
            return agora_local().date()

        return data.date()

    except Exception:
        return agora_local().date()

def formatar_mensagem_whatsapp(ocorrencias, nome):
    msg = f"""📋 RELATÓRIO DE OCORRÊNCIAS
👤 Aluno: {nome}
📅 Data do Relatório: {datetime.now().strftime('%d/%m/%y às %H:%M')}
==============================\n"""

    for i, ocorr in enumerate(ocorrencias, start=1):
        data_txt = ocorr.get("data", "")
        data_formatada = data_txt
        if data_txt:
            for fmt in ("%d-%m-%Y %H:%M:%S", "%d-%m-%Y %H:%M"):
                try:
                    data_obj = datetime.strptime(data_txt, fmt)
                    data_formatada = data_obj.strftime("%Y/%m/%d às %H:%M")
                    break
                except ValueError:
                    continue
        msg += f"""
🔸 Ocorrência {i}
📅 Data: {data_formatada}
📝 Descrição: {ocorr['descricao']}
-------------------------"""

    msg += """

👨‍🏫 Escola [CCM Pinheiros do Parana]
📞 Contato: [41 3364-6581]

Este relatório foi gerado automaticamente pelo Sistema de Ocorrências."""
    return msg

def exportar_ocorrencias_para_word(ocorrencias, nome_arquivo):
    import os
    import base64
    from io import BytesIO
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ===== BRASÃO CENTRALIZADO =====
    caminho_logo = os.path.join(os.getcwd(), "BRASÃO1.png")
    if os.path.exists(caminho_logo):
        paragrafo_logo = doc.add_paragraph()
        paragrafo_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragrafo_logo.add_run()
        run.add_picture(caminho_logo, width=Inches(1.5))

    # Título centralizado
    titulo = doc.add_heading("RELATÓRIO DE OCORRÊNCIAS", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== OCORRÊNCIAS =====
    for ocorr in ocorrencias:

        doc.add_paragraph().add_run("Aluno: ").bold = True
        doc.paragraphs[-1].add_run(f"{ocorr.get('nome','')}")

        doc.add_paragraph().add_run("CGM: ").bold = True
        doc.paragraphs[-1].add_run(f"{ocorr.get('cgm','')}")

        doc.add_paragraph().add_run("Data: ").bold = True
        doc.paragraphs[-1].add_run(f"{ocorr.get('data','')}")

        doc.add_paragraph().add_run("Descrição: ").bold = True
        doc.paragraphs[-1].add_run(f"{ocorr.get('descricao','')}")

        doc.add_paragraph("")

        # ===== ATA ANEXADA =====
        ata_base64 = ocorr.get("ata")

        if ata_base64:
            try:
                arquivo_bytes = base64.b64decode(ata_base64)

                # Se for PDF
                if arquivo_bytes[:4] == b"%PDF":
                    try:
                        from pdf2image import convert_from_bytes
                        imagens = convert_from_bytes(arquivo_bytes)

                        doc.add_paragraph("ATA Anexada:")

                        for img in imagens:
                            img_stream = BytesIO()
                            img.save(img_stream, format="PNG")
                            img_stream.seek(0)
                            doc.add_picture(img_stream, width=Inches(5))

                    except Exception:
                        doc.add_paragraph(
                            "ATA em PDF (não foi possível converter)."
                        )

                else:
                    # Se for imagem
                    img_stream = BytesIO(arquivo_bytes)
                    doc.add_paragraph("ATA Anexada:")
                    doc.add_picture(img_stream, width=Inches(5))

            except Exception:
                doc.add_paragraph("Erro ao carregar ATA.")

        doc.add_page_break()

    caminho = os.path.join(os.getcwd(), nome_arquivo)
    doc.save(caminho)

    return caminho

def exportar_ocorrencias_para_pdf(ocorrencias, nome_arquivo):
    import os
    import base64
    from io import BytesIO
    from datetime import datetime
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors

    caminho = os.path.join(os.getcwd(), nome_arquivo)
    doc = SimpleDocTemplate(caminho)

    elementos = []
    styles = getSampleStyleSheet()

    cabecalho_style = ParagraphStyle(
        'Cabecalho',
        parent=styles['Normal'],
        fontSize=12,
        textColor=colors.black,
        alignment=1,
        spaceAfter=6
    )

    normal_style = ParagraphStyle(
        'NormalCustom',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=4
    )

    # ===== BRASÃO =====
    caminho_logo = os.path.join(os.getcwd(), "BRASÃO1.png")
    if os.path.exists(caminho_logo):
        logo = Image(caminho_logo, width=1.5*inch, height=1.5*inch)
        logo.hAlign = 'CENTER'
        elementos.append(logo)
        elementos.append(Spacer(1, 15))

    # ===== CABEÇALHO =====
    elementos.append(Paragraph(
        "<b>COLÉGIO CÍVICO MILITAR PINHEIRO DO PARANA</b>",
        cabecalho_style
    ))

    elementos.append(Paragraph(
        "Relatório Oficial de Ocorrências",
        cabecalho_style
    ))

    elementos.append(Paragraph(
        f"Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}",
        cabecalho_style
    ))

    elementos.append(Spacer(1, 20))

    # ===== OCORRÊNCIAS =====
    for ocorr in ocorrencias:

        elementos.append(Paragraph(f"<b>Aluno:</b> {ocorr.get('nome','')}", normal_style))
        elementos.append(Paragraph(f"<b>CGM:</b> {ocorr.get('cgm','')}", normal_style))
        elementos.append(Paragraph(f"<b>Data:</b> {ocorr.get('data','')}", normal_style))
        elementos.append(Paragraph(f"<b>Descrição:</b> {ocorr.get('descricao','')}", normal_style))
        elementos.append(Spacer(1, 15))

        ata_base64 = ocorr.get("ata")

        if ata_base64:
            try:
                arquivo_bytes = base64.b64decode(ata_base64)

                if arquivo_bytes[:4] == b"%PDF":
                    try:
                        from pdf2image import convert_from_bytes
                        imagens = convert_from_bytes(arquivo_bytes)

                        elementos.append(Paragraph("<b>ATA Anexada:</b>", normal_style))
                        elementos.append(Spacer(1, 10))

                        for img in imagens:
                            img_stream = BytesIO()
                            img.save(img_stream, format="PNG")
                            img_stream.seek(0)

                            elementos.append(Image(img_stream, width=5*inch, height=7*inch))
                            elementos.append(Spacer(1, 10))

                    except Exception:
                        elementos.append(
                            Paragraph("ATA em PDF (não foi possível converter).", normal_style)
                        )

                else:
                    # ===== SE FOR IMAGEM =====
                    img_stream = BytesIO(arquivo_bytes)

                    elementos.append(Paragraph("<b>ATA Anexada:</b>", normal_style))
                    elementos.append(Spacer(1, 10))
                    elementos.append(Image(img_stream, width=5*inch, height=7*inch))

            except Exception:
                elementos.append(
                    Paragraph("Erro ao carregar ATA.", normal_style)
                )

        elementos.append(PageBreak())

    doc.build(elementos)

    return caminho

# --- Login ---
def pagina_login():
    st.markdown("## 👤 Login de Usuário - V2.0.3 by Leandro Malheiros")
    usuario = st.text_input("Usuário").strip()
    senha = st.text_input("Senha", type="password").strip()

    if st.button("Entrar"):
        senha_hash = hashlib.sha256(senha.encode()).hexdigest()
        user = db.usuarios.find_one({
            "usuario": usuario,
            "senha": senha_hash
        })

        if user:
            st.session_state["logado"] = True
            st.session_state["usuario"] = usuario
            st.session_state["nivel"] = user.get("nivel", "user")
            st.success("✅ Login realizado com sucesso!")
            st.rerun()
        else:
            st.error("Usuário ou senha inválidos.")

# --- Cadastro de Alunos ---
def pagina_cadastro():
    st.markdown("## ✏️ Cadastro de Alunos")

    # --- Lista de alunos cadastrados ---
    alunos = list(db.alunos.find().sort("nome", 1))

    nomes_exibicao = [""] + [
        f"{a['nome']} (CGM: {a['cgm']})"
        for a in alunos
    ]

    selecionado = st.selectbox("🔎 Buscar aluno para Alterar ou Excluir:", nomes_exibicao)

    aluno_carregado = None
    if selecionado and selecionado != "":
        # Extrai CGM do texto selecionado
        cgm_busca = selecionado.split("CGM:")[1].replace(")", "").strip()
        aluno_carregado = db.alunos.find_one({"cgm": cgm_busca})

        st.success(f"Aluno carregado: {aluno_carregado['nome']} (CGM {aluno_carregado['cgm']})")

    # --- Formulário de Cadastro ou Alteração ---
    with st.form("form_cadastro"):

        cgm = st.text_input("CGM", value=aluno_carregado["cgm"] if aluno_carregado else "")
        nome = st.text_input("Nome", value=aluno_carregado["nome"] if aluno_carregado else "")
        data = st.date_input(
        "Data de Nascimento",
            value=data_segura(aluno_carregado.get("data") if aluno_carregado else None)
    )
        telefone = st.text_input("Telefone", value=aluno_carregado["telefone"] if aluno_carregado else "")
        responsavel = st.text_input("Responsável", value=aluno_carregado["responsavel"] if aluno_carregado else "")
        turma = st.text_input("Turma", value=aluno_carregado["turma"] if aluno_carregado else "")

        col1, col2, col3 = st.columns([1,1,1])
        salvar = col1.form_submit_button("💾 Salvar / Alterar")
        excluir = col2.form_submit_button("🗑️ Excluir")
        limpar = col3.form_submit_button("🧹 Limpar")

    # --- Ações após clique ---
    if salvar:
        if cgm and nome:
            db.alunos.update_one({"cgm": cgm}, {
                "$set": {
                    "cgm": cgm,
                    "nome": nome,
                    "data": str(data),
                    "telefone": telefone,
                    "responsavel": responsavel,
                    "turma": turma
                }
            }, upsert=True)
            st.success("✅ Aluno salvo ou atualizado com sucesso!")
            st.rerun()
        else:
            st.error("Preencha todos os campos obrigatórios.")

    if excluir and aluno_carregado:
        confirmacao = st.warning(f"Tem certeza que deseja excluir o aluno {aluno_carregado['nome']} (CGM {aluno_carregado['cgm']})?")
        if st.button("✅ Confirmar Exclusão"):
            db.alunos.delete_one({"cgm": aluno_carregado["cgm"]})
            st.success("✅ Aluno excluído com sucesso!")
            st.rerun()

    if limpar:
        st.rerun()

   # --- Importação de alunos via arquivo ---
	st.subheader("📥 Importar Alunos via TXT ou CSV")
	arquivo = st.file_uploader("Escolha o arquivo .txt ou .csv", type=["txt", "csv"])

	if arquivo is not None:
		try:
			# ===== TENTA LER AUTOMÁTICO =====
			df_import = pd.read_csv(arquivo, sep=None, engine="python")

			# ===== CORREÇÃO PARA ARQUIVO BUGADO (TUDO EM 1 COLUNA) =====
			if len(df_import.columns) == 1:
				st.warning("⚠️ Arquivo veio em uma única coluna. Tentando corrigir automaticamente...")

				arquivo.seek(0)
				df_import = pd.read_csv(arquivo, delimiter="\t")

			# ===== PADRONIZA COLUNAS =====
			df_import.columns = [col.strip().lower() for col in df_import.columns]

			# ===== RENOMEAR COLUNAS POSSÍVEIS =====
			mapa_colunas = {
				"nome do estudante": "nome",
				"data de nasc.": "data",
				"telefone": "telefone",
				"turma": "turma",
				"cgm": "cgm"
			}

			df_import.rename(columns=mapa_colunas, inplace=True)

			st.success("✅ Arquivo carregado com sucesso!")
			st.dataframe(df_import)

			if st.button("🚀 Importar para o Sistema"):
				erros = []
				total_importados = 0

				for _, row in df_import.iterrows():
					try:
						cgm = str(row.get("cgm", "")).strip()
						nome = str(row.get("nome", "")).strip()
						data = str(row.get("data", "")).strip()
						telefone = str(row.get("telefone", "")).strip()
						responsavel = str(row.get("responsavel", "")).strip()
						turma = str(row.get("turma", "")).strip()

						if not cgm or not nome:
							erros.append(f"❌ CGM ou Nome ausente: {row.to_dict()}")
							continue

						aluno = {
							"cgm": cgm,
							"nome": nome,
							"data": data,
							"telefone": telefone,
							"responsavel": responsavel,
							"turma": turma
						}

						db.alunos.update_one(
							{"cgm": cgm},
							{"$set": aluno},
							upsert=True
						)

						total_importados += 1

					except Exception as e:
						erros.append(f"Erro: {row.to_dict()} → {e}")

				st.success(f"✅ Importação finalizada: {total_importados} alunos")

				if erros:
					st.warning("⚠️ Problemas encontrados:")
					for erro in erros:
						st.error(erro)

		except Exception as e:
			st.error(f"❌ Erro ao ler o arquivo: {e}")

def pagina_ocorrencias():
    st.markdown("## 🚨 Registro de Ocorrência")

    alunos = list(db.alunos.find())
    alunos_ordenados = sorted(alunos, key=lambda x: x['nome'])

    busca_cgm = st.text_input("🔍 Buscar aluno por CGM")

    if busca_cgm:
        aluno_cgm = next((a for a in alunos_ordenados if a["cgm"] == busca_cgm), None)
        if aluno_cgm:
            nomes = [f"{aluno_cgm['nome']} (CGM: {aluno_cgm['cgm']})"]
        else:
            st.warning("Nenhum aluno encontrado com esse CGM.")
            return
    else:
        nomes = [""] + [f"{a['nome']} (CGM: {a['cgm']})" for a in alunos_ordenados]

    if nomes:
        selecionado = st.selectbox("Selecione o aluno:", nomes)

        if selecionado != "":
            cgm = selecionado.split("CGM: ")[1].replace(")", "")
            nome = selecionado.split(" (CGM:")[0]

            ocorrencias = list(db.ocorrencias.find({"cgm": cgm}))
            opcoes_ocorrencias = ["Nova Ocorrência"] + [
                f"{o['data']} - {o['descricao'][:30]}..." for o in ocorrencias
            ]

            ocorrencia_selecionada = st.selectbox("📌 Ocorrência:", opcoes_ocorrencias)

            descricao = ""
            ata = ""

            # ================= NOVA OCORRÊNCIA =================
            if ocorrencia_selecionada == "Nova Ocorrência":
                descricao = st.text_area("✏️ Descrição da Ocorrência", key="descricao_nova")
                ata = st.text_input("📄 ATA (opcional)", key="ata_nova")

                arquivo_ata = st.file_uploader(
                    "📤 Importar ATA (Somente JPG)",
                    type=["jpg", "jpeg"],
                    key="upload_ata_nova"
                )

                if arquivo_ata:
                    ata = base64.b64encode(arquivo_ata.read()).decode("utf-8")

                if st.button("✅ Registrar Nova Ocorrência", key="btn_nova") and descricao:
                    agora = agora_local().strftime("%Y-%m-%d %H:%M:%S")
                    telefone = next((a['telefone'] for a in alunos if a['cgm'] == cgm), "")

                    db.ocorrencias.insert_one({
                        "cgm": cgm,
                        "nome": nome,
                        "telefone": telefone,
                        "data": agora,
                        "descricao": descricao,
                        "ata": ata
                    })

                    st.success("✅ Ocorrência registrada com sucesso!")

            # ================= OCORRÊNCIA EXISTENTE =================
            else:
                index = opcoes_ocorrencias.index(ocorrencia_selecionada) - 1
                ocorrencia = ocorrencias[index]

                descricao = st.text_area(
                    "✏️ Descrição da Ocorrência",
                    value=ocorrencia.get("descricao", ""),
                    key=f"desc_{ocorrencia['_id']}"
                )

                ata = st.text_input(
                    "📄 ATA (opcional)",
                    value=ocorrencia.get("ata", ""),
                    key=f"ata_{ocorrencia['_id']}"
                )

                arquivo_ata = st.file_uploader(
                    "📤 Importar nova ATA (Somente em JPG)",
                    type=["jpg", "jpeg"],
                    key=f"upload_ata_{ocorrencia['_id']}"
                )

                if arquivo_ata:
                    ata = base64.b64encode(arquivo_ata.read()).decode("utf-8")

                col1, col2 = st.columns(2)

                with col1:
                    if st.button("💾 Alterar Ocorrência", key=f"alt_{ocorrencia['_id']}"):
                        db.ocorrencias.update_one(
                            {"_id": ocorrencia["_id"]},
                            {"$set": {
                                "descricao": descricao,
                                "ata": ata
                            }}
                        )
                        st.success("✅ Ocorrência atualizada com sucesso!")

                with col2:
                    confirmar_exclusao = st.checkbox(
                        "Confirmar exclusão",
                        key=f"conf_{ocorrencia['_id']}"
                    )
                    if confirmar_exclusao:
                        if st.button("🗑️ Excluir Ocorrência", key=f"del_{ocorrencia['_id']}"):
                            db.ocorrencias.delete_one({"_id": ocorrencia["_id"]})
                            st.success("🗑️ Ocorrência excluída com sucesso!")
                            st.rerun()

def pagina_exportar():
    import urllib
    import uuid

    st.markdown("## 📥 Exportar Relatórios")

    resultados = list(db.ocorrencias.find({}))
    if not resultados:
        st.warning("Nenhuma ocorrência encontrada.")
        return

    # ===================== BUSCA POR CGM =====================
    st.subheader("🔍 Buscar por CGM")
    cgm_input = st.text_input("Digite o CGM do aluno")
    col1, col2 = st.columns(2)

    if col1.button("📄 Gerar Word por CGM", key="word_cgm") and cgm_input:
        dados = list(db.ocorrencias.find({"cgm": cgm_input}))
        if dados:
            caminho = exportar_ocorrencias_para_word(
                dados, f"ocorrencias_{cgm_input}.docx"
            )
            with open(caminho, "rb") as f:
                st.download_button(
                    "📥 Baixar Word",
                    f.read(),
                    file_name=f"ocorrencias_{cgm_input}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    if col2.button("🧾 Gerar PDF por CGM", key="pdf_cgm") and cgm_input:
        dados = list(db.ocorrencias.find({"cgm": cgm_input}))
        if dados:
            caminho = exportar_ocorrencias_para_pdf(
                dados, f"ocorrencias_{cgm_input}.pdf"
            )
            with open(caminho, "rb") as f:
                st.download_button(
                    "📥 Baixar PDF",
                    f.read(),
                    file_name=f"ocorrencias_{cgm_input}.pdf",
                    mime="application/pdf"
                )

    # ===================== PERÍODO =====================
    st.subheader("📅 Exportar por Período")

    uid = str(uuid.uuid4())

    data_inicio = st.date_input("Data inicial", key=f"ini_{uid}")
    data_fim = st.date_input("Data final", key=f"fim_{uid}")

    if st.button("🔎 Gerar relatório por período", key=f"periodo_{uid}"):

        if data_fim < data_inicio:
            st.error("A data final não pode ser menor que a data inicial.")
            st.stop()

        inicio = data_inicio.strftime("%Y-%m-%d 00:00:00")
        fim = data_fim.strftime("%Y-%m-%d 23:59:59")

        dados = list(db.ocorrencias.find({
            "data": {
                "$gte": inicio,
                "$lte": fim
            }
        }))

        if not dados:
            st.warning("Nenhuma ocorrência encontrada no período selecionado.")
        else:
            st.success(f"{len(dados)} ocorrência(s) encontrada(s).")

            caminho_docx = exportar_ocorrencias_para_word(
                dados,
                "relatorio_periodo.docx"
            )

            with open(caminho_docx, "rb") as f:
                st.download_button(
                    "📥 Baixar DOCX",
                    f.read(),
                    file_name="relatorio_periodo.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_docx_{uid}"
                )

            caminho_pdf = exportar_ocorrencias_para_pdf(
                dados,
                "relatorio_periodo.pdf"
            )

            with open(caminho_pdf, "rb") as f:
                st.download_button(
                    "📥 Baixar PDF",
                    f.read(),
                    file_name="relatorio_periodo.pdf",
                    mime="application/pdf",
                    key=f"download_pdf_{uid}"
                )

    # ===================== AGRUPADO POR ALUNO =====================
    st.subheader("📄 Relatórios Individuais por Aluno")

    ocorrencias_por_aluno = {}

    for ocorr in resultados:
        nome = ocorr.get("nome", "")
        ocorrencias_por_aluno.setdefault(nome, []).append(ocorr)

    for nome, lista in sorted(ocorrencias_por_aluno.items()):

        with st.expander(f"📄 Relatório de {nome}"):

            telefone = lista[0].get("telefone", "")

            for ocorr in lista:
                st.write(f"📅 {ocorr.get('data', '')} - 📝 {ocorr.get('descricao', '')}")

            mensagem = formatar_mensagem_whatsapp(lista, nome)

            st.text_area(
                "📋 WhatsApp",
                mensagem,
                height=200,
                key=f"msg_{nome}_{lista[0]['_id']}"
            )

            if telefone:
                numero = (
                    telefone.replace("(", "")
                    .replace(")", "")
                    .replace("-", "")
                    .replace(" ", "")
                )

                link = f"https://api.whatsapp.com/send?phone=55{numero}&text={urllib.parse.quote(mensagem)}"

                st.markdown(f"[📱 Enviar para {telefone}]({link})")

            col1, col2 = st.columns(2)

            if col1.button("📄 Gerar DOCX", key=f"doc_{nome}_{lista[0]['_id']}"):
                caminho = exportar_ocorrencias_para_word(
                    lista,
                    f"relatorio_{nome.replace(' ','_')}.docx"
                )

                with open(caminho, "rb") as f:
                    st.download_button(
                        "📥 Baixar DOCX",
                        f.read(),
                        file_name=f"relatorio_{nome.replace(' ','_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            if col2.button("🧾 Gerar PDF", key=f"pdf_{nome}_{lista[0]['_id']}"):
                caminho = exportar_ocorrencias_para_pdf(
                    lista,
                    f"relatorio_{nome.replace(' ','_')}.pdf"
                )

                with open(caminho, "rb") as f:
                    st.download_button(
                        "📥 Baixar PDF",
                        f.read(),
                        file_name=f"relatorio_{nome.replace(' ','_')}.pdf",
                        mime="application/pdf"
                    )

# --- Lista de Alunos ---
def pagina_lista():
    st.markdown("## 📄 Lista de Alunos")
    dados = list(db.alunos.find({}, {"_id": 0}))
    if dados:
        df = pd.DataFrame(dados)
        st.dataframe(df.sort_values("nome"))
    else:
        st.info("Nenhum aluno cadastrado.")

# --- Cadastro de Usuários ---
def pagina_usuarios():
    st.markdown("## 👥 Cadastro de Usuários")
    
    # Exemplo de segurança: só admin pode cadastrar
    if st.session_state.get("nivel") != "admin":
        st.warning("Apenas administradores podem cadastrar novos usuários.")
        return

    # Formulário de cadastro
    with st.form("form_usuarios"):
        usuario = st.text_input("Novo usuário")
        senha = st.text_input("Senha", type="password")
        nivel = st.selectbox("Nível de acesso", ["user", "admin"])
        cadastrar = st.form_submit_button("Cadastrar")

    if cadastrar:
        usuario = usuario.strip()
        senha = senha.strip()
        if usuario and senha:
            senha_hash = hashlib.sha256(senha.encode()).hexdigest()
            try:
                resultado = db.usuarios.insert_one({
                    "usuario": usuario,
                    "senha": senha_hash,
                    "nivel": nivel
                })
                st.success("✅ Usuário cadastrado com sucesso!")
                print("Usuário salvo com id:", resultado.inserted_id)
            except Exception as e:
                print("Erro ao salvar usuário:", e)
                st.error(f"Erro ao salvar usuário: {e}")
        else:
            st.error("Preencha todos os campos.")

    if st.button("👀 Ver Usuários Salvos"):
        usuarios = list(db.usuarios.find())
        if usuarios:
            for u in usuarios:
                st.write(u)
        else:
            st.info("Nenhum usuário cadastrado ainda.")

# --- Menu Lateral ---
def menu():
    st.sidebar.image("BRASÃO1.png", use_container_width=True)
    st.sidebar.markdown("### 📚 Menu de Navegação")
    opcoes = ["Cadastro", "Ocorrências", "Exportar", "Lista"]
    if st.session_state.get("nivel") == "admin":
        opcoes.append("Usuários")
    pagina = st.sidebar.selectbox("Escolha a aba:", opcoes)

    if pagina == "Cadastro":
        pagina_cadastro()
    elif pagina == "Ocorrências":
        pagina_ocorrencias()
    elif pagina == "Exportar":
        pagina_exportar()
    elif pagina == "Lista":
        pagina_lista()
    elif pagina == "Usuários":
        pagina_usuarios()

def sair():
    st.session_state["logado"] = False
    st.session_state["usuario"] = ""
    st.session_state["nivel"] = ""
    st.rerun()
    
# --- Execução ---
if "logado" not in st.session_state:
    st.session_state["logado"] = False

if not st.session_state["logado"]:
    pagina_login()
else:
    if st.sidebar.button("🚪 Sair do Sistema"):
        sair()
    else:
        menu()
